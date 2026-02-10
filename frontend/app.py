import streamlit as st
import requests
from docx import Document
from io import BytesIO

# -------------------------------------------------
# Helper: Create Question Paper Word file
# -------------------------------------------------
def create_question_file(questions, question_type):
    doc = Document()
    doc.add_heading("Question Paper", level=1)

    for i, q in enumerate(questions, 1):
        if question_type == "mcq":
            doc.add_paragraph(f"Q{i}. {q['question']}")
            options = ["A", "B", "C", "D"]
            for idx, opt in enumerate(q["options"]):
                doc.add_paragraph(f"{options[idx]}. {opt}")
            doc.add_paragraph("")
        else:
            doc.add_paragraph(f"Q{i}. {q['question']}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -------------------------------------------------
# Helper: Create Answer Key Word file
# -------------------------------------------------
def create_answer_file(questions, question_type):
    doc = Document()
    doc.add_heading("Answer Key", level=1)

    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"Q{i}. {q['question']}")

        if question_type == "mcq":
            doc.add_paragraph(f"Correct Answer: {q['correct_answer']}")
        else:
            doc.add_paragraph(f"Model Answer:\n{q['model_answer']}")
            doc.add_paragraph("Key Points:")
            for kp in q["key_points"]:
                doc.add_paragraph(f"- {kp}")
            doc.add_paragraph(
                "Expected Keywords: " + ", ".join(q["expected_keywords"])
            )

        doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



TEXT_API = "http://127.0.0.1:8000/generate-questions"
PDF_API = "http://127.0.0.1:8000/generate-questions-from-pdf"

st.set_page_config("Question Generator", "📘")
st.title("📘 Intelligent Question Generation System")

tab1, tab2 = st.tabs(["📝 Text / Topic", "📄 PDF"])

# =================================================
# TEXT / TOPIC
# =================================================
with tab1:
    mode = st.radio(
        "Input type",
        ["Topic based", "Content based"],
        key="text_mode"
    )

    topic = None
    content = None

    if mode == "Topic based":
        topic = st.text_input("Topic", key="text_topic")
    else:
        content = st.text_area("Content", height=200, key="text_content")

    keywords_input = st.text_input(
        "Keywords (optional, comma-separated)",
        key="text_keywords"
    )

    keywords = (
        [k.strip() for k in keywords_input.split(",") if k.strip()]
        if keywords_input else None
    )

    num_questions = st.number_input(
        "Number of Questions",
        1, 20, 5,
        key="text_num_questions"
    )

    difficulty = st.selectbox(
        "Difficulty",
        ["easy", "medium", "hard"],
        key="text_difficulty"
    )

    question_type = st.radio(
        "Question Type",
        ["descriptive", "mcq"],
        horizontal=True,
        key="text_question_type"
    )

    if st.button("Generate Questions", key="text_generate_btn"):
        payload = {
            "topic": topic,
            "content": content,
            "keywords": keywords,
            "num_questions": num_questions,
            "difficulty": difficulty,
            "question_type": question_type
        }

        res = requests.post(TEXT_API, json=payload)
        if res.status_code == 200:
            st.session_state["questions"] = res.json()["questions"]
            st.session_state["question_type"] = question_type
        else:
            st.error(res.text)


# =================================================
# PDF
# =================================================
with tab2:
    uploaded_pdf = st.file_uploader(
        "Upload PDF",
        type=["pdf"],
        key="pdf_uploader"
    )

    pdf_num_questions = st.number_input(
        "Number of Questions",
        1, 50, 10,
        key="pdf_num_questions"
    )

    pdf_difficulty = st.selectbox(
        "Difficulty",
        ["easy", "medium", "hard"],
        key="pdf_difficulty"
    )

    pdf_question_type = st.radio(
        "Question Type",
        ["descriptive", "mcq"],
        horizontal=True,
        key="pdf_question_type"
    )

    if uploaded_pdf and st.button("Generate Questions", key="pdf_generate_btn"):
        files = {
            "file": (uploaded_pdf.name, uploaded_pdf.getvalue(), "application/pdf")
        }

        params = {
            "num_questions": pdf_num_questions,
            "difficulty": pdf_difficulty,
            "question_type": pdf_question_type
        }

        res = requests.post(PDF_API, files=files, params=params)
        if res.status_code == 200:
            st.session_state["questions"] = res.json()["questions"]
            st.session_state["question_type"] = pdf_question_type
        else:
            st.error(res.text)


# =================================================
# OUTPUT
# =================================================
if "questions" in st.session_state:
    st.subheader("Generated Questions")

    qtype = st.session_state.get("question_type", "descriptive")

    for i, q in enumerate(st.session_state["questions"], 1):
        if qtype == "mcq":
            st.write(f"**Q{i}. {q['question']}**")
            options = ["A", "B", "C", "D"]
            for idx, opt in enumerate(q["options"]):
                st.write(f"{options[idx]}. {opt}")
        else:
            st.write(f"**Q{i}. {q['question']}**")

    col1, col2 = st.columns(2)

    with col1:
        st.download_button(
            "📄 Download Question Paper",
            create_question_file(
                st.session_state["questions"],
                qtype
            ),
            "generated_questions.docx"
        )

    with col2:
        st.download_button(
            "📘 Download Answers",
            create_answer_file(
                st.session_state["questions"],
                qtype
            ),
            "answer_key.docx"
        )
